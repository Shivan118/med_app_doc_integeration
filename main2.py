import streamlit as st
import google.generativeai as genai
from google_api_key import google_api_key
from PIL import Image
import io
import docx
from io import BytesIO

# Configure Google Generative AI
genai.configure(api_key=google_api_key)

# Set up the model
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 0,
    "max_output_tokens": 8192,
}

safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
]

system_prompts = {
    "image": """
    You are a domain expert in medical image analysis. You are tasked with 
    examining medical images for a renowned hospital. Your expertise will help in identifying 
    or discovering any anomalies, diseases, conditions or health issues that might be present in the image.
    
    Your key responsibilities:
    1. Detailed Analysis: Scrutinize and thoroughly examine each image, 
       focusing on finding any abnormalities.
    2. Analysis Report: Document all the findings and 
       clearly articulate them in a structured format.
    3. Recommendations: Based on the analysis, suggest remedies, 
       tests or treatments as applicable.
    4. Treatments: If applicable, lay out detailed treatments 
       which can help in faster recovery.
    
    Important Notes to remember:
    1. Scope of response: Only respond if the image pertains to human health issues.
    2. Clarity of image: In case the image is unclear, 
       note that certain aspects are 'Unable to be correctly determined based on the uploaded image'
    3. Disclaimer: Accompany your analysis with the disclaimer: 
       "This is an AI BOT made by MEDI360. Consult with a Doctor before making any decisions."
    
    Please provide the final response in headings and sub-headings in bullet format: 
    Detailed Analysis, Analysis Report, Recommendations and Treatments.

    when you found \n or \n\n in responses extend the output in next line.

    Note: If images are not related to medical topics, as you're a Medical AI Chatbot, 
    please inform the user that you can only analyze medical-related images.
    """,

    "text": """
    You are an AI medical assistant designed to provide concise, accurate information on various medical topics. Your role is to offer brief, helpful responses to health-related queries, including suggestions for medications when appropriate.

    When responding to queries, structure your answer concisely with these key areas as applicable:

    1. Brief Overview: Concise information about the medical topic or condition.
    2. Key Symptoms: List 3-5 main symptoms if relevant.
    3. Medication Suggestions: 
       - Suggest 2-3 commonly used medications for the condition, including generic names.
       - Briefly explain the purpose of each medication.
       - Include typical dosage forms (e.g., tablet, syrup, injection).
    4. Quick Suggestions: Offer 2-3 practical health tips or lifestyle modifications.
    5. When to Seek Help: Briefly state when professional medical attention is necessary.

    Important guidelines:
    1. Keep responses short and to the point, typically within 200-250 words.
    2. Use simple, clear language accessible to a general audience.
    3. When suggesting medications, always preface with: "Common medications that a doctor might consider include: "
    4. For all responses, end with this disclaimer: "Note: This is AI-generated information by MEDI360. Always consult a healthcare professional for personalized medical advice and before taking any medication."

    Provide your response in a simple format with appropriate headings.

    Remember, you're an AI assistant providing general information, not a replacement for professional medical advice or diagnosis. Always encourage users to consult with a healthcare provider for specific medical advice and treatment.
    """
}

# Update the model initialization to use the new prompts
model_image = genai.GenerativeModel(model_name="gemini-1.5-pro-latest",
                                    generation_config=generation_config,
                                    safety_settings=safety_settings)

model_text = genai.GenerativeModel(model_name="gemini-1.5-pro-latest",
                                   generation_config=generation_config,
                                   safety_settings=safety_settings)

def is_xray(image):
    # This is a simple check. You might want to implement a more sophisticated method
    # to detect X-ray images, possibly using machine learning.
    return image.mode == 'L'  # Check if the image is grayscale

def generate_doc(content, is_xray_image=False):
    doc = docx.Document()
    
    # Split the content into sections
    sections = content.split('**')
    
    for i, section in enumerate(sections):
        if i % 2 == 0:  # Even indexes are normal text or empty
            continue
        else:  # Odd indexes are headers or content
            # Add the header in bold
            p = doc.add_paragraph()
            p.add_run(section.strip()).bold = True
            
            # If there's content following this header, add it as bullet points
            if i + 1 < len(sections):
                content = sections[i + 1]
                bullet_points = content.split('*')
                for point in bullet_points:
                    if point.strip():
                        doc.add_paragraph(point.strip(), style='List Bullet')
    
    # Add the disclaimer as a separate paragraph
    #disclaimer = "This is an AI BOT made by MEDI360. Consult with a Doctor before making any decisions."
    #doc.add_paragraph(disclaimer)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

st.title("Visual Medical Assistant 👨‍⚕️ 🩺 🏥")
st.subheader("An app to help with medical analysis using images and text")

tab1, tab2 = st.tabs(["Image Analysis", "Text Query"])

with tab1:
    st.header("Image Analysis")
    uploaded_file = st.file_uploader("Upload the image for Analysis:", type=["jpg", "jpeg", "png"])
    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption="Uploaded Image", use_column_width=True)
        if st.button("Analyze"):
            with st.spinner("Analyzing..."):
                image_data = uploaded_file.getvalue()
                image_parts = [{"mime_type": "image/jpeg", "data": image_data}]
                prompt_parts = [image_parts[0], system_prompts["image"]]
                ai_response = model_image.generate_content(prompt_parts)
                if ai_response:
                    st.markdown(ai_response.text)

                    
                    is_xray_image = is_xray(image)
                    
                    # Generate and offer download of .doc file
                    doc_io = generate_doc(ai_response.text)
                    st.download_button(
                        label="Download analysis as .doc",
                        data=doc_io,
                        file_name="medical_image_analysis.doc",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("An error occurred during analysis.")

with tab2:
    st.header("Text Query")
    query = st.text_area("Enter your medical query:")
    if st.button("Submit Query"):
        with st.spinner("Processing..."):
            prompt_parts = [
                system_prompts["text"],
                f"User query: {query}"
            ]
            ai_response = model_text.generate_content(prompt_parts)
            if ai_response:
                st.markdown(ai_response.text)
                
                # Generate and offer download of .doc file
                doc_io = generate_doc(ai_response.text)
                st.download_button(
                    label="Download response as .doc",
                    data=doc_io,
                    file_name="medical_query_response.doc",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("An error occurred while processing your query.")

st.sidebar.title("About")
st.sidebar.info("This is a Visual Medical Assistant app that uses AI to analyze medical images and answer medical queries.")